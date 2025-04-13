import os
import base64
from fpdf import FPDF
from docx import Document
import tempfile
from pathlib import Path
import streamlit as st


class DocumentExporter:
    """Handles exporting documents in various formats"""

    @staticmethod
    def export_to_pdf(title, original_text, simplified_text, translated_text=None, language=None):
        """
        Export document content to PDF with Unicode support using fpdf2
        """

        FONT_PATH = Path(__file__).resolve().parent.parent / "assets" / "ArialUnicode.ttf"

        # Input validation
        if title is None:
            title = "Untitled"
        if original_text is None:
            original_text = ""
        if simplified_text is None:
            simplified_text = ""

        # Create a PDF with fpdf2
        pdf = FPDF(orientation='P', unit='mm', format='A4')
        # Set margins
        pdf.set_margins(20, 20, 20)
        pdf.add_page()
        
        pdf.add_font("ArialUnicode", "",str(FONT_PATH) , uni=True)
        
        # With fpdf2, we can directly work with Unicode text
        # No need for the sanitize_text function
        
        # Process text blocks with better fpdf2 features
        def process_text_block(text, header):
            # Add header with fpdf2 positioning
            pdf.set_font("ArialUnicode", size=12)
            pdf.cell(w=0, h=10, text=header, new_x="LMARGIN", new_y="NEXT")
            pdf.set_font("ArialUnicode", size=10)
            
            # Safety check
            if text is None:
                return
                
            # With fpdf2, we can use multi_cell more effectively
            paragraphs = str(text).split('\n')
            for paragraph in paragraphs:
                if paragraph.strip():  # Skip empty paragraphs
                    pdf.multi_cell(w=0, h=5, text=paragraph)
                    pdf.ln(2)  # Small space after paragraph
        
        # Document header
        pdf.set_font("ArialUnicode", size=16)
        pdf.cell(w=0, h=10, text="Legal Document Simplification", 
                 align="C", new_x="LMARGIN", new_y="NEXT")
        
        # Document title
        pdf.set_font("ArialUnicode", size=12)
        safe_title = str(title)[:40] if title else "Untitled"
        pdf.cell(w=0, h=10, text=f"Document: {safe_title}", 
                 new_x="LMARGIN", new_y="NEXT")
        pdf.ln(5)  # Space after title
        
        # Process each text block
        process_text_block(original_text, "Original Text:")
        pdf.ln(5)
        process_text_block(simplified_text, "Simplified Text:")
        
        # Handle translated text
        if translated_text and language:
            pdf.ln(5)
            process_text_block(translated_text, f"Translated Text ({language}):")
        
        # Add timestamp
        pdf.ln(5)
        pdf.set_font("ArialUnicode", size=8)
        timestamp = str(st.session_state.get('timestamp', 'N/A'))
        pdf.cell(w=0, h=5, text=f"Generated on: {timestamp}", 
                 new_x="LMARGIN", new_y="NEXT")
        
        # With fpdf2, we can directly return bytes
        return pdf.output(dest='bytes').decode('latin-1')

    @staticmethod
    def export_to_docx(title, original_text, simplified_text, translated_text=None, language=None):
        """
        Export document content to DOCX (Word) format
        """
        # Input validation
        if title is None:
            title = "Untitled"
        if original_text is None:
            original_text = ""
        if simplified_text is None:
            simplified_text = ""

        # Convert all inputs to strings
        title = str(title)
        original_text = str(original_text)
        simplified_text = str(simplified_text)
        if translated_text is not None:
            translated_text = str(translated_text)
        if language is not None:
            language = str(language)

        doc = Document()

        # Add title
        doc.add_heading("Legal Document Simplification", 0)
        doc.add_heading(f"Document: {title[:50]}", 1)

        # Original Text
        doc.add_heading("Original Text:", 2)
        doc.add_paragraph(original_text)

        # Simplified Text
        doc.add_heading("Simplified Text:", 2)
        doc.add_paragraph(simplified_text)

        # Translated Text (if available)
        if translated_text and language:
            doc.add_heading(f"Translated Text ({language}):", 2)
            doc.add_paragraph(translated_text)

        # Add timestamp
        timestamp = str(st.session_state.get('timestamp', 'N/A'))
        doc.add_paragraph(f"Generated on: {timestamp}").italic = True

        # Save to a BytesIO object
        with tempfile.NamedTemporaryFile(delete=False, suffix='.docx') as tmp:
            doc.save(tmp.name)
            tmp.seek(0)
            with open(tmp.name, "rb") as f:
                docx_bytes = f.read()

        # Clean up the temp file
        os.unlink(tmp.name)

        return docx_bytes

    @staticmethod
    def export_to_txt(title, original_text, simplified_text, translated_text=None, language=None):
        """
        Export document content to plain text
        """
        # Input validation
        if title is None:
            title = "Untitled"
        if original_text is None:
            original_text = ""
        if simplified_text is None:
            simplified_text = ""

        # Convert all inputs to strings
        title = str(title)
        original_text = str(original_text)
        simplified_text = str(simplified_text)
        if translated_text is not None:
            translated_text = str(translated_text)
        if language is not None:
            language = str(language)

        content = []
        content.append("LEGAL DOCUMENT SIMPLIFICATION")
        content.append(f"Document: {title}\n")

        content.append("ORIGINAL TEXT:")
        content.append(original_text)
        content.append("")

        content.append("SIMPLIFIED TEXT:")
        content.append(simplified_text)
        content.append("")

        if translated_text and language:
            content.append(f"TRANSLATED TEXT ({language}):")
            content.append(translated_text)
            content.append("")

        timestamp = str(st.session_state.get('timestamp', 'N/A'))
        content.append(f"Generated on: {timestamp}")

        # Join with newlines and convert to bytes
        return "\n".join(content).encode("utf-8")

    @staticmethod
    def get_download_link(file_bytes, filename, file_format, display_text):
        """
        Generate a download link for a file
        """
        # Ensure filename is valid
        filename = str(filename) if filename else "Simplified_document" 
        filename = filename.replace(" ", "_")
        
        # Ensure we have bytes (different encoding for PDF vs text)
        if not isinstance(file_bytes, bytes):
            if file_format == "pdf":
                file_bytes = str(file_bytes).encode("latin-1")  # PDF binary data needs latin-1
            else:
                file_bytes = str(file_bytes).encode("utf-8")  # Text data can use utf-8
                
        # Rest of the method remains the same
        b64 = base64.b64encode(file_bytes).decode()
        
        # Map format to MIME type
        mime_types = {
            "pdf": "application/pdf",
            "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            "txt": "text/plain"
        }
        mime_type = mime_types.get(file_format, "application/octet-stream")
        
        href = f'<a href="data:{mime_type};base64,{b64}" download="{filename}.{file_format}">{display_text}</a>'
        return href

    @staticmethod
    def render_export_options(title, original_text, simplified_text, translated_text=None, language=None):
        """Render export options in the Streamlit UI"""
        st.markdown("### Export Document")

        # Create filename base
        filename_base = title.replace(
            " ", "_")[:30] if title else "Legal_Document"

        # Export format selection
        export_format = st.selectbox(
            "Select Format:",
            ["", "PDF", "Word Document", "Text File"],
            key="export_format"
        )

        # Only show download button if format is selected
        if export_format:
            if export_format == "PDF":
                pdf_bytes = DocumentExporter.export_to_pdf(
                    title, original_text, simplified_text, translated_text, language
                )
                download_link = DocumentExporter.get_download_link(
                    pdf_bytes, filename_base, "pdf", "Download PDF"
                )
                st.markdown(download_link, unsafe_allow_html=True)

            elif export_format == "Word Document":
                docx_bytes = DocumentExporter.export_to_docx(
                    title, original_text, simplified_text, translated_text, language
                )
                download_link = DocumentExporter.get_download_link(
                    docx_bytes, filename_base, "docx", "Download Word"
                )
                st.markdown(download_link, unsafe_allow_html=True)

            elif export_format == "Text File":
                txt_bytes = DocumentExporter.export_to_txt(
                    title, original_text, simplified_text, translated_text, language
                )
                download_link = DocumentExporter.get_download_link(
                    txt_bytes, filename_base, "txt", "Download Text"
                )
                st.markdown(download_link, unsafe_allow_html=True)
